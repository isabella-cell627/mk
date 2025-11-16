from flask import Flask, render_template, request, jsonify, send_file
import os
import json
from datetime import datetime
import markdown
from weasyprint import HTML, CSS
from io import BytesIO
import secrets
from werkzeug.utils import secure_filename
from flask_sqlalchemy import SQLAlchemy
from sqlalchemy.orm import DeclarativeBase
from werkzeug.middleware.proxy_fix import ProxyFix

class Base(DeclarativeBase):
    pass

db = SQLAlchemy(model_class=Base)

app = Flask(__name__)
app.secret_key = os.environ.get("SESSION_SECRET")
app.wsgi_app = ProxyFix(app.wsgi_app, x_proto=1, x_host=1)

app.config["SQLALCHEMY_DATABASE_URI"] = os.environ.get("DATABASE_URL")
app.config["SQLALCHEMY_ENGINE_OPTIONS"] = {
    "pool_recycle": 300,
    "pool_pre_ping": True,
}

db.init_app(app)

UPLOADS_DIR = 'uploads'
EXPORTS_DIR = 'exports'

os.makedirs(UPLOADS_DIR, exist_ok=True)
os.makedirs(EXPORTS_DIR, exist_ok=True)

def init_db():
    with app.app_context():
        import models
        db.create_all()

def safe_join_path(base_dir, filename):
    """Safely join a filename to a base directory, preventing path traversal attacks."""
    safe_name = secure_filename(filename)
    if not safe_name:
        raise ValueError("Invalid filename")
    
    full_path = os.path.abspath(os.path.join(base_dir, safe_name))
    base_path = os.path.abspath(base_dir)
    
    if not full_path.startswith(base_path + os.sep):
        raise ValueError("Path traversal attempt detected")
    
    return full_path, safe_name

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/api/save', methods=['POST'])
def save_file():
    from models import Document
    try:
        data = request.get_json(silent=True)
        if not data:
            return jsonify({'success': False, 'error': 'No data provided'}), 400
        
        filename = data.get('filename', 'untitled.md')
        content = data.get('content', '')
        folder_id = data.get('folder_id')
        document_id = data.get('document_id')
        
        if not filename.endswith('.md'):
            filename += '.md'
        
        if document_id:
            document = Document.query.get(document_id)
            if document:
                document.filename = filename
                document.content = content
                if folder_id is not None:
                    document.folder_id = folder_id
                document.updated_at = datetime.utcnow()
            else:
                document = Document(
                    filename=filename,
                    content=content,
                    folder_id=folder_id
                )
                db.session.add(document)
        else:
            query = Document.query.filter_by(filename=filename)
            if folder_id is not None:
                query = query.filter_by(folder_id=folder_id)
            else:
                query = query.filter_by(folder_id=None)
            
            existing = query.first()
            if existing:
                existing.content = content
                existing.updated_at = datetime.utcnow()
                document = existing
            else:
                document = Document(
                    filename=filename,
                    content=content,
                    folder_id=folder_id
                )
                db.session.add(document)
        
        db.session.commit()
        
        return jsonify({
            'success': True,
            'message': 'File saved successfully',
            'filename': filename,
            'document_id': document.id,
            'document': document.to_dict(include_content=True)
        })
    except ValueError as e:
        db.session.rollback()
        return jsonify({'success': False, 'error': str(e)}), 400
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/open', methods=['POST'])
def open_file():
    from models import Document, RecentFile
    try:
        data = request.get_json(silent=True)
        if not data:
            return jsonify({'success': False, 'error': 'No data provided'}), 400
        
        filename = data.get('filename')
        document_id = data.get('document_id')
        
        if document_id:
            document = Document.query.get(document_id)
            if not document:
                return jsonify({'success': False, 'error': 'Document not found'}), 404
        elif filename:
            document = Document.query.filter_by(filename=filename).first()
            if not document:
                return jsonify({'success': False, 'error': 'Document not found'}), 404
        else:
            return jsonify({'success': False, 'error': 'No filename or document_id provided'}), 400
        
        document.last_opened_at = datetime.utcnow()
        
        recent = RecentFile(document_id=document.id, accessed_at=datetime.utcnow())
        db.session.add(recent)
        db.session.commit()
        
        return jsonify({
            'success': True,
            'content': document.content,
            'filename': document.filename,
            'document_id': document.id,
            'document': document.to_dict(include_content=True)
        })
    except ValueError as e:
        db.session.rollback()
        return jsonify({'success': False, 'error': str(e)}), 400
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/files', methods=['GET'])
def list_files():
    from models import Document
    try:
        documents = Document.query.order_by(Document.updated_at.desc()).all()
        files = [{
            'filename': doc.filename,
            'size': len(doc.content.encode('utf-8')) if doc.content else 0,
            'modified': doc.updated_at.isoformat() if doc.updated_at else None,
            'document_id': doc.id,
            'folder_id': doc.folder_id,
            'is_favorite': doc.is_favorite,
            'is_pinned': doc.is_pinned
        } for doc in documents]
        
        return jsonify({'success': True, 'files': files})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/export/html', methods=['POST'])
def export_html():
    try:
        data = request.get_json(silent=True)
        if not data:
            return jsonify({'success': False, 'error': 'No data provided'}), 400
        
        content = data.get('content', '')
        filename = data.get('filename', 'document')
        
        html_content = markdown.markdown(
            content,
            extensions=['extra', 'codehilite', 'tables', 'fenced_code']
        )
        
        safe_title = secure_filename(filename.replace('.md', ''))
        
        full_html = f"""
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{safe_title}</title>
    <style>
        body {{
            font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, Oxygen, Ubuntu, Cantarell, sans-serif;
            max-width: 800px;
            margin: 0 auto;
            padding: 2rem;
            line-height: 1.6;
            color: #333;
        }}
        h1, h2, h3, h4, h5, h6 {{
            margin-top: 1.5em;
            margin-bottom: 0.5em;
        }}
        code {{
            background: #f4f4f4;
            padding: 0.2em 0.4em;
            border-radius: 3px;
            font-family: 'Courier New', monospace;
        }}
        pre {{
            background: #f4f4f4;
            padding: 1em;
            border-radius: 5px;
            overflow-x: auto;
        }}
        pre code {{
            background: none;
            padding: 0;
        }}
        blockquote {{
            border-left: 4px solid #ddd;
            margin-left: 0;
            padding-left: 1em;
            color: #666;
        }}
        table {{
            border-collapse: collapse;
            width: 100%;
            margin: 1em 0;
        }}
        th, td {{
            border: 1px solid #ddd;
            padding: 0.5em;
            text-align: left;
        }}
        th {{
            background: #f4f4f4;
        }}
    </style>
</head>
<body>
{html_content}
</body>
</html>
"""
        
        export_filename = f"{safe_title}_export.html"
        export_path, _ = safe_join_path(EXPORTS_DIR, export_filename)
        
        with open(export_path, 'w', encoding='utf-8') as f:
            f.write(full_html)
        
        return send_file(
            export_path,
            as_attachment=True,
            download_name=export_filename,
            mimetype='text/html'
        )
    except ValueError as e:
        return jsonify({'success': False, 'error': str(e)}), 400
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/export/pdf', methods=['POST'])
def export_pdf():
    try:
        data = request.get_json(silent=True)
        if not data:
            return jsonify({'success': False, 'error': 'No data provided'}), 400
        
        content = data.get('content', '')
        filename = data.get('filename', 'document')
        
        html_content = markdown.markdown(
            content,
            extensions=['extra', 'codehilite', 'tables', 'fenced_code']
        )
        
        full_html = f"""
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <style>
        @page {{
            size: A4;
            margin: 2cm;
        }}
        body {{
            font-family: 'DejaVu Sans', Arial, sans-serif;
            line-height: 1.6;
            color: #333;
        }}
        h1, h2, h3, h4, h5, h6 {{
            margin-top: 1.5em;
            margin-bottom: 0.5em;
            page-break-after: avoid;
        }}
        code {{
            background: #f4f4f4;
            padding: 0.2em 0.4em;
            border-radius: 3px;
            font-family: 'Courier New', monospace;
        }}
        pre {{
            background: #f4f4f4;
            padding: 1em;
            border-radius: 5px;
            overflow-x: auto;
            page-break-inside: avoid;
        }}
        pre code {{
            background: none;
            padding: 0;
        }}
        blockquote {{
            border-left: 4px solid #ddd;
            margin-left: 0;
            padding-left: 1em;
            color: #666;
        }}
        table {{
            border-collapse: collapse;
            width: 100%;
            margin: 1em 0;
            page-break-inside: avoid;
        }}
        th, td {{
            border: 1px solid #ddd;
            padding: 0.5em;
            text-align: left;
        }}
        th {{
            background: #f4f4f4;
        }}
    </style>
</head>
<body>
{html_content}
</body>
</html>
"""
        
        pdf_buffer = BytesIO()
        HTML(string=full_html).write_pdf(pdf_buffer)
        pdf_buffer.seek(0)
        
        safe_filename = secure_filename(filename.replace('.md', ''))
        export_filename = f"{safe_filename}_export.pdf"
        
        return send_file(
            pdf_buffer,
            as_attachment=True,
            download_name=export_filename,
            mimetype='application/pdf'
        )
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/export/txt', methods=['POST'])
def export_txt():
    try:
        data = request.get_json(silent=True)
        if not data:
            return jsonify({'success': False, 'error': 'No data provided'}), 400
        
        content = data.get('content', '')
        filename = data.get('filename', 'document')
        
        safe_filename = secure_filename(filename.replace('.md', ''))
        export_filename = f"{safe_filename}_export.txt"
        export_path, _ = safe_join_path(EXPORTS_DIR, export_filename)
        
        with open(export_path, 'w', encoding='utf-8') as f:
            f.write(content)
        
        return send_file(
            export_path,
            as_attachment=True,
            download_name=export_filename,
            mimetype='text/plain'
        )
    except ValueError as e:
        return jsonify({'success': False, 'error': str(e)}), 400
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/folders', methods=['GET'])
def get_folders():
    from models import Folder
    try:
        root_folders = Folder.query.filter_by(parent_id=None).order_by(Folder.position).all()
        return jsonify({
            'success': True,
            'folders': [folder.to_dict() for folder in root_folders]
        })
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/folders', methods=['POST'])
def create_folder():
    from models import Folder
    try:
        data = request.get_json()
        folder = Folder(
            name=data.get('name', 'New Folder'),
            parent_id=data.get('parent_id'),
            color=data.get('color', '#6366f1'),
            icon=data.get('icon', 'folder')
        )
        db.session.add(folder)
        db.session.commit()
        return jsonify({'success': True, 'folder': folder.to_dict()})
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/folders/<int:folder_id>', methods=['PUT'])
def update_folder(folder_id):
    from models import Folder
    try:
        folder = Folder.query.get_or_404(folder_id)
        data = request.get_json()
        
        if 'name' in data:
            folder.name = data['name']
        if 'parent_id' in data:
            folder.parent_id = data['parent_id']
        if 'color' in data:
            folder.color = data['color']
        if 'icon' in data:
            folder.icon = data['icon']
        if 'position' in data:
            folder.position = data['position']
        
        db.session.commit()
        return jsonify({'success': True, 'folder': folder.to_dict()})
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/folders/<int:folder_id>', methods=['DELETE'])
def delete_folder(folder_id):
    from models import Folder
    try:
        folder = Folder.query.get_or_404(folder_id)
        db.session.delete(folder)
        db.session.commit()
        return jsonify({'success': True})
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/documents', methods=['GET'])
def get_documents():
    from models import Document
    try:
        folder_id = request.args.get('folder_id', type=int)
        
        if folder_id is not None:
            documents = Document.query.filter_by(folder_id=folder_id).order_by(Document.updated_at.desc()).all()
        else:
            documents = Document.query.order_by(Document.updated_at.desc()).all()
        
        return jsonify({
            'success': True,
            'documents': [doc.to_dict() for doc in documents]
        })
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/documents', methods=['POST'])
def create_document():
    from models import Document
    try:
        data = request.get_json()
        document = Document(
            filename=data.get('filename', 'untitled.md'),
            content=data.get('content', ''),
            folder_id=data.get('folder_id')
        )
        db.session.add(document)
        db.session.commit()
        return jsonify({'success': True, 'document': document.to_dict(include_content=True)})
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/documents/<int:doc_id>', methods=['GET'])
def get_document(doc_id):
    from models import Document, RecentFile
    try:
        document = Document.query.get_or_404(doc_id)
        document.last_opened_at = datetime.utcnow()
        
        recent = RecentFile(document_id=doc_id, accessed_at=datetime.utcnow())
        db.session.add(recent)
        db.session.commit()
        
        return jsonify({'success': True, 'document': document.to_dict(include_content=True)})
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/documents/<int:doc_id>', methods=['PUT'])
def update_document(doc_id):
    from models import Document
    try:
        document = Document.query.get_or_404(doc_id)
        data = request.get_json()
        
        if 'filename' in data:
            document.filename = data['filename']
        if 'content' in data:
            document.content = data['content']
        if 'folder_id' in data:
            document.folder_id = data['folder_id']
        if 'is_favorite' in data:
            document.is_favorite = data['is_favorite']
        if 'is_pinned' in data:
            document.is_pinned = data['is_pinned']
        
        document.updated_at = datetime.utcnow()
        db.session.commit()
        return jsonify({'success': True, 'document': document.to_dict(include_content=True)})
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/documents/<int:doc_id>', methods=['DELETE'])
def delete_document(doc_id):
    from models import Document
    try:
        document = Document.query.get_or_404(doc_id)
        db.session.delete(document)
        db.session.commit()
        return jsonify({'success': True})
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/tags', methods=['GET'])
def get_tags():
    from models import Tag
    try:
        tags = Tag.query.order_by(Tag.name).all()
        return jsonify({
            'success': True,
            'tags': [tag.to_dict() for tag in tags]
        })
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/tags', methods=['POST'])
def create_tag():
    from models import Tag
    try:
        data = request.get_json()
        tag = Tag(
            name=data.get('name'),
            color=data.get('color', '#6366f1')
        )
        db.session.add(tag)
        db.session.commit()
        return jsonify({'success': True, 'tag': tag.to_dict()})
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/tags/<int:tag_id>', methods=['DELETE'])
def delete_tag(tag_id):
    from models import Tag
    try:
        tag = Tag.query.get_or_404(tag_id)
        db.session.delete(tag)
        db.session.commit()
        return jsonify({'success': True})
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/documents/<int:doc_id>/tags', methods=['POST'])
def add_tag_to_document(doc_id):
    from models import Document, Tag
    try:
        document = Document.query.get_or_404(doc_id)
        data = request.get_json()
        tag_id = data.get('tag_id')
        
        tag = Tag.query.get_or_404(tag_id)
        if tag not in document.tags:
            document.tags.append(tag)
            db.session.commit()
        
        return jsonify({'success': True, 'document': document.to_dict()})
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/documents/<int:doc_id>/tags/<int:tag_id>', methods=['DELETE'])
def remove_tag_from_document(doc_id, tag_id):
    from models import Document, Tag
    try:
        document = Document.query.get_or_404(doc_id)
        tag = Tag.query.get_or_404(tag_id)
        
        if tag in document.tags:
            document.tags.remove(tag)
            db.session.commit()
        
        return jsonify({'success': True, 'document': document.to_dict()})
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/categories', methods=['GET'])
def get_categories():
    from models import Category
    try:
        categories = Category.query.order_by(Category.name).all()
        return jsonify({
            'success': True,
            'categories': [cat.to_dict() for cat in categories]
        })
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/categories', methods=['POST'])
def create_category():
    from models import Category
    try:
        data = request.get_json()
        category = Category(
            name=data.get('name'),
            color=data.get('color', '#ec4899'),
            icon=data.get('icon', 'bookmark')
        )
        db.session.add(category)
        db.session.commit()
        return jsonify({'success': True, 'category': category.to_dict()})
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/categories/<int:cat_id>', methods=['DELETE'])
def delete_category(cat_id):
    from models import Category
    try:
        category = Category.query.get_or_404(cat_id)
        db.session.delete(category)
        db.session.commit()
        return jsonify({'success': True})
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/documents/<int:doc_id>/categories', methods=['POST'])
def add_category_to_document(doc_id):
    from models import Document, Category
    try:
        document = Document.query.get_or_404(doc_id)
        data = request.get_json()
        cat_id = data.get('category_id')
        
        category = Category.query.get_or_404(cat_id)
        if category not in document.categories:
            document.categories.append(category)
            db.session.commit()
        
        return jsonify({'success': True, 'document': document.to_dict()})
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/documents/<int:doc_id>/categories/<int:cat_id>', methods=['DELETE'])
def remove_category_from_document(doc_id, cat_id):
    from models import Document, Category
    try:
        document = Document.query.get_or_404(doc_id)
        category = Category.query.get_or_404(cat_id)
        
        if category in document.categories:
            document.categories.remove(category)
            db.session.commit()
        
        return jsonify({'success': True, 'document': document.to_dict()})
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/search', methods=['GET'])
def search_documents():
    from models import Document, Tag, Category
    try:
        query = request.args.get('q', '').strip()
        tag_ids = request.args.getlist('tags', type=int)
        cat_ids = request.args.getlist('categories', type=int)
        is_favorite = request.args.get('favorite', type=lambda v: v.lower() == 'true')
        is_pinned = request.args.get('pinned', type=lambda v: v.lower() == 'true')
        
        docs_query = Document.query
        
        if query:
            docs_query = docs_query.filter(
                db.or_(
                    Document.filename.ilike(f'%{query}%'),
                    Document.content.ilike(f'%{query}%')
                )
            )
        
        if tag_ids:
            for tag_id in tag_ids:
                docs_query = docs_query.filter(Document.tags.any(Tag.id == tag_id))
        
        if cat_ids:
            for cat_id in cat_ids:
                docs_query = docs_query.filter(Document.categories.any(Category.id == cat_id))
        
        if is_favorite is not None:
            docs_query = docs_query.filter_by(is_favorite=is_favorite)
        
        if is_pinned is not None:
            docs_query = docs_query.filter_by(is_pinned=is_pinned)
        
        documents = docs_query.order_by(Document.updated_at.desc()).all()
        
        return jsonify({
            'success': True,
            'documents': [doc.to_dict() for doc in documents],
            'count': len(documents)
        })
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/recent', methods=['GET'])
def get_recent_files():
    from models import RecentFile
    try:
        limit = request.args.get('limit', 10, type=int)
        recent_files = RecentFile.query.order_by(RecentFile.accessed_at.desc()).limit(limit).all()
        
        return jsonify({
            'success': True,
            'recent': [rf.to_dict() for rf in recent_files if rf.document]
        })
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/favorites', methods=['GET'])
def get_favorites():
    from models import Document
    try:
        documents = Document.query.filter_by(is_favorite=True).order_by(Document.updated_at.desc()).all()
        return jsonify({
            'success': True,
            'documents': [doc.to_dict() for doc in documents]
        })
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/pinned', methods=['GET'])
def get_pinned():
    from models import Document
    try:
        documents = Document.query.filter_by(is_pinned=True).order_by(Document.updated_at.desc()).all()
        return jsonify({
            'success': True,
            'documents': [doc.to_dict() for doc in documents]
        })
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/migrate-legacy-files', methods=['POST'])
def migrate_legacy_files():
    from models import Document
    try:
        if not os.path.exists(UPLOADS_DIR):
            return jsonify({'success': True, 'message': 'No legacy files directory found', 'migrated': 0})
        
        migrated_count = 0
        errors = []
        
        for filename in os.listdir(UPLOADS_DIR):
            if not filename.endswith('.md'):
                continue
            
            try:
                filepath = os.path.join(UPLOADS_DIR, filename)
                
                existing = Document.query.filter_by(filename=filename, folder_id=None).first()
                if existing:
                    continue
                
                with open(filepath, 'r', encoding='utf-8') as f:
                    content = f.read()
                
                stat = os.stat(filepath)
                created_time = datetime.fromtimestamp(stat.st_ctime)
                modified_time = datetime.fromtimestamp(stat.st_mtime)
                
                document = Document(
                    filename=filename,
                    content=content,
                    folder_id=None,
                    created_at=created_time,
                    updated_at=modified_time
                )
                db.session.add(document)
                migrated_count += 1
            except Exception as e:
                errors.append(f"{filename}: {str(e)}")
        
        db.session.commit()
        
        return jsonify({
            'success': True,
            'message': f'Migrated {migrated_count} files',
            'migrated': migrated_count,
            'errors': errors
        })
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/export/docx', methods=['POST'])
def export_docx():
    from docx import Document as DocxDocument
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import re
    
    try:
        data = request.get_json(silent=True)
        if not data:
            return jsonify({'success': False, 'error': 'No data provided'}), 400
        
        content = data.get('content', '')
        filename = data.get('filename', 'document')
        
        doc = DocxDocument()
        
        lines = content.split('\n')
        i = 0
        while i < len(lines):
            line = lines[i].rstrip()
            
            if line.startswith('# '):
                heading = doc.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                heading = doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                heading = doc.add_heading(line[4:], level=3)
            elif line.startswith('#### '):
                heading = doc.add_heading(line[5:], level=4)
            elif line.startswith('- ') or line.startswith('* '):
                p = doc.add_paragraph(line[2:], style='List Bullet')
            elif re.match(r'^\d+\. ', line):
                p = doc.add_paragraph(re.sub(r'^\d+\. ', '', line), style='List Number')
            elif line.startswith('> '):
                p = doc.add_paragraph(line[2:])
                p.style = 'Quote'
            elif line.startswith('```'):
                code_block = []
                i += 1
                while i < len(lines) and not lines[i].startswith('```'):
                    code_block.append(lines[i])
                    i += 1
                if code_block:
                    p = doc.add_paragraph('\n'.join(code_block))
                    p.style = 'No Spacing'
                    for run in p.runs:
                        run.font.name = 'Courier New'
                        run.font.size = Pt(10)
            elif line.strip():
                p = doc.add_paragraph(line)
            else:
                doc.add_paragraph()
            
            i += 1
        
        safe_filename = secure_filename(filename.replace('.md', ''))
        export_filename = f"{safe_filename}_export.docx"
        export_path, _ = safe_join_path(EXPORTS_DIR, export_filename)
        
        doc.save(export_path)
        
        return send_file(
            export_path,
            as_attachment=True,
            download_name=export_filename,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/export/odt', methods=['POST'])
def export_odt():
    from odf.opendocument import OpenDocumentText
    from odf.text import P, H
    
    try:
        data = request.get_json(silent=True)
        if not data:
            return jsonify({'success': False, 'error': 'No data provided'}), 400
        
        content = data.get('content', '')
        filename = data.get('filename', 'document')
        
        textdoc = OpenDocumentText()
        
        lines = content.split('\n')
        for line in lines:
            line = line.rstrip()
            
            if line.startswith('# '):
                h = H(outlinelevel=1, text=line[2:])
                textdoc.text.addElement(h)
            elif line.startswith('## '):
                h = H(outlinelevel=2, text=line[3:])
                textdoc.text.addElement(h)
            elif line.startswith('### '):
                h = H(outlinelevel=3, text=line[4:])
                textdoc.text.addElement(h)
            elif line.strip():
                p = P(text=line)
                textdoc.text.addElement(p)
            else:
                textdoc.text.addElement(P())
        
        safe_filename = secure_filename(filename.replace('.md', ''))
        export_filename = f"{safe_filename}_export.odt"
        export_path, _ = safe_join_path(EXPORTS_DIR, export_filename)
        
        textdoc.save(export_path)
        
        return send_file(
            export_path,
            as_attachment=True,
            download_name=export_filename,
            mimetype='application/vnd.oasis.opendocument.text'
        )
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/export/latex', methods=['POST'])
def export_latex():
    import re
    
    try:
        data = request.get_json(silent=True)
        if not data:
            return jsonify({'success': False, 'error': 'No data provided'}), 400
        
        content = data.get('content', '')
        filename = data.get('filename', 'document')
        
        latex_content = []
        latex_content.append(r'\documentclass{article}')
        latex_content.append(r'\usepackage[utf8]{inputenc}')
        latex_content.append(r'\usepackage{listings}')
        latex_content.append(r'\usepackage{xcolor}')
        latex_content.append(r'\usepackage{hyperref}')
        latex_content.append(r'\begin{document}')
        latex_content.append('')
        
        lines = content.split('\n')
        in_code_block = False
        code_lang = ''
        
        for line in lines:
            if line.startswith('```'):
                if not in_code_block:
                    in_code_block = True
                    code_lang = line[3:].strip()
                    latex_content.append(r'\begin{lstlisting}')
                else:
                    latex_content.append(r'\end{lstlisting}')
                    in_code_block = False
            elif in_code_block:
                latex_content.append(line)
            elif line.startswith('# '):
                latex_content.append(r'\section{' + line[2:].replace('_', r'\_').replace('#', r'\#') + '}')
            elif line.startswith('## '):
                latex_content.append(r'\subsection{' + line[3:].replace('_', r'\_').replace('#', r'\#') + '}')
            elif line.startswith('### '):
                latex_content.append(r'\subsubsection{' + line[4:].replace('_', r'\_').replace('#', r'\#') + '}')
            elif line.startswith('- ') or line.startswith('* '):
                if not latex_content[-1].startswith(r'\item'):
                    latex_content.append(r'\begin{itemize}')
                latex_content.append(r'\item ' + line[2:])
                if lines.index(line) + 1 >= len(lines) or not (lines[lines.index(line) + 1].startswith('- ') or lines[lines.index(line) + 1].startswith('* ')):
                    latex_content.append(r'\end{itemize}')
            elif line.strip():
                escaped_line = line.replace('_', r'\_').replace('#', r'\#').replace('&', r'\&').replace('%', r'\%')
                escaped_line = re.sub(r'\*\*(.+?)\*\*', r'\\textbf{\1}', escaped_line)
                escaped_line = re.sub(r'\*(.+?)\*', r'\\textit{\1}', escaped_line)
                escaped_line = re.sub(r'`(.+?)`', r'\\texttt{\1}', escaped_line)
                latex_content.append(escaped_line)
            else:
                latex_content.append('')
        
        latex_content.append('')
        latex_content.append(r'\end{document}')
        
        safe_filename = secure_filename(filename.replace('.md', ''))
        export_filename = f"{safe_filename}_export.tex"
        export_path, _ = safe_join_path(EXPORTS_DIR, export_filename)
        
        with open(export_path, 'w', encoding='utf-8') as f:
            f.write('\n'.join(latex_content))
        
        return send_file(
            export_path,
            as_attachment=True,
            download_name=export_filename,
            mimetype='application/x-latex'
        )
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/export/epub', methods=['POST'])
def export_epub():
    from ebooklib import epub
    import markdown as md
    
    try:
        data = request.get_json(silent=True)
        if not data:
            return jsonify({'success': False, 'error': 'No data provided'}), 400
        
        content = data.get('content', '')
        filename = data.get('filename', 'document')
        
        book = epub.EpubBook()
        
        safe_title = secure_filename(filename.replace('.md', ''))
        book.set_identifier(f'id_{safe_title}')
        book.set_title(safe_title)
        book.set_language('en')
        book.add_author('MarkPro Editor')
        
        html_content = md.markdown(
            content,
            extensions=['extra', 'codehilite', 'tables', 'fenced_code']
        )
        
        c1 = epub.EpubHtml(title=safe_title, file_name='chap_01.xhtml', lang='en')
        c1.content = f'<html><head></head><body>{html_content}</body></html>'
        
        book.add_item(c1)
        
        book.toc = [epub.Link('chap_01.xhtml', safe_title, 'intro')]
        
        book.add_item(epub.EpubNcx())
        book.add_item(epub.EpubNav())
        
        book.spine = ['nav', c1]
        
        safe_filename = secure_filename(filename.replace('.md', ''))
        export_filename = f"{safe_filename}_export.epub"
        export_path, _ = safe_join_path(EXPORTS_DIR, export_filename)
        
        epub.write_epub(export_path, book, {})
        
        return send_file(
            export_path,
            as_attachment=True,
            download_name=export_filename,
            mimetype='application/epub+zip'
        )
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/export/presentation', methods=['POST'])
def export_presentation():
    try:
        data = request.get_json(silent=True)
        if not data:
            return jsonify({'success': False, 'error': 'No data provided'}), 400
        
        content = data.get('content', '')
        filename = data.get('filename', 'document')
        
        slides = content.split('\n---\n')
        
        html_slides = []
        for slide in slides:
            slide_html = markdown.markdown(
                slide.strip(),
                extensions=['extra', 'codehilite', 'tables', 'fenced_code']
            )
            html_slides.append(f'<section>{slide_html}</section>')
        
        safe_title = secure_filename(filename.replace('.md', ''))
        
        full_html = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{safe_title} - Presentation</title>
    <link rel="stylesheet" href="https://cdnjs.cloudflare.com/ajax/libs/reveal.js/4.5.0/reset.min.css">
    <link rel="stylesheet" href="https://cdnjs.cloudflare.com/ajax/libs/reveal.js/4.5.0/reveal.min.css">
    <link rel="stylesheet" href="https://cdnjs.cloudflare.com/ajax/libs/reveal.js/4.5.0/theme/black.min.css">
    <link rel="stylesheet" href="https://cdnjs.cloudflare.com/ajax/libs/reveal.js/4.5.0/plugin/highlight/monokai.min.css">
    <style>
        .reveal h1, .reveal h2, .reveal h3, .reveal h4, .reveal h5, .reveal h6 {{
            text-transform: none;
        }}
        .reveal pre code {{
            max-height: 500px;
        }}
    </style>
</head>
<body>
    <div class="reveal">
        <div class="slides">
            {chr(10).join(html_slides)}
        </div>
    </div>
    
    <script src="https://cdnjs.cloudflare.com/ajax/libs/reveal.js/4.5.0/reveal.min.js"></script>
    <script src="https://cdnjs.cloudflare.com/ajax/libs/reveal.js/4.5.0/plugin/notes/notes.min.js"></script>
    <script src="https://cdnjs.cloudflare.com/ajax/libs/reveal.js/4.5.0/plugin/markdown/markdown.min.js"></script>
    <script src="https://cdnjs.cloudflare.com/ajax/libs/reveal.js/4.5.0/plugin/highlight/highlight.min.js"></script>
    <script>
        Reveal.initialize({{
            hash: true,
            plugins: [ RevealMarkdown, RevealHighlight, RevealNotes ]
        }});
    </script>
</body>
</html>"""
        
        export_filename = f"{safe_title}_presentation.html"
        export_path, _ = safe_join_path(EXPORTS_DIR, export_filename)
        
        with open(export_path, 'w', encoding='utf-8') as f:
            f.write(full_html)
        
        return send_file(
            export_path,
            as_attachment=True,
            download_name=export_filename,
            mimetype='text/html'
        )
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000, debug=True)
